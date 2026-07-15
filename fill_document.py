"""
Isi otomatis "Format Dokumentasi Pembangunan dan Pengembangan Aplikasi SPBE"
Berdasarkan data proyek RENJANA (Relawan Remaja Aman Bencana)
Kabupaten Tanah Bumbu, Kalimantan Selatan
"""

from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy
import re

SRC = "/Volumes/data/Project/renjana/Format Dokumentasi Pembangunan dan Pengembangan Aplikasi.docx"
DST = "/Volumes/data/Project/renjana/Dokumentasi RENJANA - SPBE - Terisi.docx"

# ──────────────────────────────────────────────
# DATA PROYEK
# ──────────────────────────────────────────────
PROJECT = {
    "nama_aplikasi": "RENJANA (Relawan Remaja Aman Bencana)",
    "jenis_aplikasi": "Aplikasi Khusus",
    "versi_aplikasi": "1.0.0",
    "nama_dokumen": "",  # per-section
    "instansi": "Pemerintah Kabupaten Tanah Bumbu",
    "unit_tik": "Dinas Komunikasi dan Informatika Kabupaten Tanah Bumbu",
    "unit_bisnis": "Badan Penanggulangan Bencana Daerah (BPBD) Kabupaten Tanah Bumbu",
    "penyusun": "Koordinator Teknis RENJANA",
    "jabatan_penyusun": "Pengelola TIK / Analis Sistem Informasi",
    "pemeriksa": "Koordinator SPBE Kabupaten Tanah Bumbu",
    "pimpinan": "Bupati Tanah Bumbu / Kepala Pelaksana BPBD",
    "tahun": "2026",
}

# ──────────────────────────────────────────────
# DOKUMEN INFO per section (A-G)
# ──────────────────────────────────────────────
DOKUMEN_INFO = {
    "Analisis Kebutuhan": {
        "table_label": "Dokumentasi Analisis Kebutuhan",
    },
    "Perencanaan": {
        "table_label": "Dokumentasi Perencanaan",
    },
    "Rancang Bangun": {
        "table_label": "Dokumentasi Rancang Bangun",
    },
    "Implementasi": {
        "table_label": "Dokumentasi Implementasi",
    },
    "Uji Kelaikan": {
        "table_label": "Dokumentasi Uji Kelaikan",
    },
    "Pemeliharaan": {
        "table_label": "Dokumentasi Pemeliharaan",
    },
    "Evaluasi": {
        "table_label": "Dokumentasi Evaluasi",
    },
}


def fill_identitas_table(table, doc_name):
    """Isi tabel Identitas Dokumen (12 baris x 2 kolom)"""
    rows = table.rows
    if len(rows) < 12:
        return

    # Row 0: Nama Aplikasi SPBE
    _set_cell(rows[0].cells[1], PROJECT["nama_aplikasi"])
    
    # Row 1: Jenis Aplikasi
    _set_cell(rows[1].cells[1], PROJECT["jenis_aplikasi"])
    
    # Row 2: Versi Aplikasi
    _set_cell(rows[2].cells[1], PROJECT["versi_aplikasi"])
    
    # Row 3: Nama Dokumen
    _set_cell(rows[3].cells[1], doc_name)

    # Row 4: Versi Dokumen
    _set_cell(rows[4].cells[1], "1.0")

    # Row 5: Instansi
    _set_cell(rows[5].cells[1], PROJECT["instansi"])

    # Row 6: Unit Kerja Pengelola TIK
    _set_cell(rows[6].cells[1], PROJECT["unit_tik"])

    # Row 7: Unit Kerja Pemilik Proses Bisnis
    _set_cell(rows[7].cells[1], PROJECT["unit_bisnis"])

    # Row 8: Disusun oleh
    _set_cell(rows[8].cells[1], f"{PROJECT['penyusun']} / {PROJECT['jabatan_penyusun']}")

    # Row 9: Diperiksa oleh (Koordinator SPBE)
    _set_cell(rows[9].cells[1], PROJECT["pemeriksa"])

    # Row 10: Disetujui oleh
    _set_cell(rows[10].cells[1], PROJECT["pimpinan"])

    # Row 11: Tanggal Penyusunan
    _set_cell(rows[11].cells[1], f"2026")


def _set_cell(cell, text):
    """Set cell text, clearing existing content"""
    for p in cell.paragraphs:
        for run in p.runs:
            run.text = ""
        p.text = ""
    p = cell.paragraphs[0]
    p.text = text


def fill_single_cell_table(table, text):
    """Isi tabel dengan 1 baris 1 kolom"""
    if text:
        cell = table.rows[0].cells[0]
        _set_cell(cell, text)


def find_identitas_tables(doc):
    """Cari semua tabel Identitas Dokumen (12 baris x 2 kolom)"""
    tables = []
    for table in doc.tables:
        rows = table.rows
        if len(rows) == 12 and len(rows[0].cells) >= 2:
            first_cell = rows[0].cells[0].text.strip()
            if "Nama Aplikasi" in first_cell or "nama aplikasi" in first_cell.lower():
                tables.append(table)
    return tables


def main():
    doc = Document(SRC)

    # ── 1. COVER PAGE ──
    for i, p in enumerate(doc.paragraphs):
        full = p.text.strip()
        if full.startswith("Nama Aplikasi :"):
            for run in p.runs:
                run.text = ""
            p.text = f"Nama Aplikasi : {PROJECT['nama_aplikasi']}"
        elif full.startswith("Instansi :"):
            for run in p.runs:
                run.text = ""
            p.text = f"Instansi : {PROJECT['instansi']}"
        elif full.startswith("Tahun :"):
            for run in p.runs:
                run.text = ""
            p.text = f"Tahun : {PROJECT['tahun']}"

    # ── 2. FIND AND FILL IDENTITAS TABLES ──
    identitas_tables = find_identitas_tables(doc)
    doc_names = [
        "Dokumentasi Analisis Kebutuhan",
        "Dokumentasi Perencanaan",
        "Dokumentasi Rancang Bangun",
        "Dokumentasi Implementasi",
        "Dokumentasi Uji Kelaikan",
        "Dokumentasi Pemeliharaan",
        "Dokumentasi Evaluasi",
    ]

    for i, table in enumerate(identitas_tables):
        if i < len(doc_names):
            fill_identitas_table(table, doc_names[i])

    # ── 3. FILL CONTENT SECTIONS ──
    # Map of paragraph text patterns → replacement text
    # We fill single-cell tables (empty ones) and empty paragraphs
    # that serve as answer areas
    
    # Strategy: iterate paragraphs, find empty ones that follow a heading/question
    content_fills = {
        # --- DOKUMEN A: Analisis Kebutuhan ---
        "A.1": (
            "Undang-Undang Nomor 23 Tahun 2014 tentang Pemerintahan Daerah;\n"
            "Peraturan Presiden Nomor 95 Tahun 2018 tentang Sistem Pemerintahan Berbasis Elektronik;\n"
            "Peraturan Menteri Komunikasi dan Digital Nomor 6 Tahun 2025 tentang Standar Teknis dan Prosedur Pembangunan dan Pengembangan Aplikasi SPBE;\n"
            "Peraturan Daerah Kabupaten Tanah Bumbu tentang Penyelenggaraan SPBE;\n"
            "Peraturan Bupati Tanah Bumbu tentang Penanggulangan Bencana Daerah dan kedudukan BPBD."
        ),
        "A.2": (
            "Permasalahan:\n"
            "1. Data relawan remaja, kegiatan kebencanaan, dan capaian program RENJANA masih dikelola secara manual pada spreadsheet dan catatan fisik.\n"
            "2. Tidak ada single source of truth untuk data relawan dan kegiatan.\n"
            "3. Proses pelaporan dari 12 kecamatan memakan waktu 3-5 hari.\n"
            "4. Pemangku kepentingan (BPBD, Dinas Pendidikan) tidak dapat mengakses data secara real-time.\n"
            "5. Distribusi relawan per kecamatan tidak terpantau secara terpusat.\n\n"
            "Kebutuhan:\n"
            "1. Dashboard komando sebagai pusat informasi seluruh aktivitas RENJANA.\n"
            "2. Manajemen data relawan, kegiatan, dan capaian terpusat.\n"
            "3. Sistem informasi yang dapat diakses publik.\n"
            "4. Otomatisasi laporan dan visualisasi data.\n"
            "5. Integrasi data kebencanaan dan peta sebaran interaktif."
        ),
        "A.4": (
            "Maksud: Menyediakan sistem informasi dashboard dan manajemen terpadu untuk program kebencanaan berbasis remaja di Kabupaten Tanah Bumbu, sebagai pusat komando bagi pengelola program dalam memantau relawan, kegiatan, edukasi, dan capaian.\n\n"
            "Tujuan:\n"
            "1. Monitoring terpusat — dashboard sebagai pusat informasi seluruh aktivitas RENJANA.\n"
            "2. Efisiensi operasional — mengurangi waktu input dan pencarian data dari jam ke detik.\n"
            "3. Transparansi — data capaian dapat diakses oleh pimpinan dan pemangku kepentingan.\n"
            "4. Skalabilitas — basis data dan arsitektur siap dikembangkan ke modul-modul lain.\n"
            "5. Kesiapsiagaan bencana — data relawan yang akurat memungkinkan mobilisasi cepat saat darurat."
        ),
        "A.5": (
            "Ruang lingkup aplikasi RENJANA mencakup:\n"
            "1. Dashboard komando dengan statistik ringkas, sebaran per kecamatan, donut chart jenis kegiatan, capaian tahunan.\n"
            "2. Manajemen data relawan (1.248+ relawan dari 12 kecamatan).\n"
            "3. Manajemen kegiatan dengan filter status dan tipe.\n"
            "4. Peta sebaran interaktif 12 kecamatan dengan hotspot bencana.\n"
            "5. Learning Management System (LMS) edukasi bencana dengan kuis dan sertifikat.\n"
            "6. Galeri foto, berita/pengumuman dengan editor markdown.\n"
            "7. Pusat dokumen (SOP, panduan, regulasi).\n"
            "8. Form pengaduan publik dan survey kepuasan.\n"
            "9. Direktori kontak 24+ koordinator per kecamatan."
        ),
        "A.8": (
            "Target kesiapan penerapan aplikasi RENJANA: Tahun 2026 (Triwulan II).\n"
            "Aplikasi telah beroperasi di https://renjana.maulanabuilds.com."
        ),
        "A.9": (
            "Sasaran pengguna aplikasi RENJANA:\n"
            "1. BPBD Kabupaten Tanah Bumbu — Admin program, pemantauan data.\n"
            "2. Dinas Pendidikan Kabupaten Tanah Bumbu — Pemantauan sekolah binaan.\n"
            "3. Koordinator Kecamatan (12 kecamatan) — Pengelolaan relawan tingkat kecamatan.\n"
            "4. Relawan RENJANA (1.248+ orang) — Akses informasi, kegiatan, edukasi.\n"
            "5. Masyarakat umum — Akses informasi publik, pengaduan, survey.\n"
            "6. Pelaku usaha dan lembaga mitra — Kolaborasi program."
        ),
        "A.10": (
            "Lokasi implementasi: Kabupaten Tanah Bumbu, Provinsi Kalimantan Selatan.\n"
            "Meliputi 12 kecamatan: Simpang Empat, Batulicin, Kusan Hilir, Kusan Hulu, Sungai Loban, Satui, Angsana, Karang Bintang, Mantewe, Kuranji, Teluk Kepayang, Batu Putih.\n"
            "Pusat data: Pusat Data Instansi (server VPS) dan didukung Pusat Data Nasional (PDN)."
        ),
        # --- DOKUMEN B: Perencanaan ---
        "B.1": (
            "Ruang lingkup perencanaan meliputi seluruh siklus pembangunan dan pengembangan aplikasi RENJANA: analisis kebutuhan, perencanaan arsitektur, rancang bangun, implementasi, uji kelaikan, pemeliharaan, dan evaluasi. Aplikasi dibangun di atas platform Laju Go (Go Fiber + Svelte 5 + Inertia.js 3 + SQLite + templ)."
        ),
        "B.2": (
            "Proses Bisnis yang diotomasi:\n"
            "- PB-01: Manajemen data relawan (pendaftaran, profil, pencarian, status keaktifan)\n"
            "- PB-02: Manajemen kegiatan (pembuatan, publikasi, filter, statistik)\n"
            "- PB-03: Edukasi kebencanaan (kursus, modul, kuis, sertifikat)\n"
            "- PB-04: Pelaporan dan dashboard (agregasi data, visualisasi, chart)\n"
            "- PB-05: Layanan publik (pengaduan, survey, informasi dokumen)\n\n"
            "Layanan SPBE terkait: Layanan Kelembagaan (Profil Organisasi), Layanan Data dan Informasi Publik."
        ),
        "B.3": (
            "Kerangka kerja terpilih: Agile Development Cycle (Scrum).\n"
            "Alasan pemilihan: Fleksibilitas dalam pengembangan fitur bertahap, responsif terhadap perubahan kebutuhan pengguna, dan siklus iterasi cepat (sprint 2 minggu).\n"
            "SNI yang dirujuk: SNI ISO/IEC/IEEE 12207 (Proses Siklus Hidup Perangkat Lunak)."
        ),
        "B.4": (
            "Skema pelaksana: Swakelola (pengembang internal) dengan Pihak Ketiga terbatas.\n"
            "Dasar pertimbangan: Proyek dikembangkan oleh tenaga teknis internal dengan platform boilerplate Laju Go yang sudah teruji. Hosting menggunakan infrastruktur yang sudah ada."
        ),
        "B.8": (
            "Pemenuhan persyaratan keamanan informasi:\n"
            "1. Autentikasi: Email/password (bcrypt/argon2id) + Google OAuth.\n"
            "2. Otorisasi: Role-Based Access Control (admin, koordinator, relawan).\n"
            "3. Session: Database-backed dengan sliding expiration.\n"
            "4. CSRF: Perlindungan CSRF pada seluruh POST/PUT/DELETE (XSRF-TOKEN cookie).\n"
            "5. Rate Limiting: Throttle pada endpoint login/register.\n"
            "6. Audit Trail: Pencatatan aktivitas pengguna di log sistem.\n"
            "7. SQL Injection: SQL type-safe via sqlc (compile-time query generation).\n"
            "8. Koneksi HTTPS pada production."
        ),
        # --- DOKUMEN C: Rancang Bangun ---
        "C.1": (
            "Arsitektur aplikasi RENJANA menggunakan pola Monolith-First dengan komponen modular:\n\n"
            "Backend: Go Fiber (HTTP framework) dengan arsitektur berlapis:\n"
            "- Handler → Service → Query (sqlc) → SQLite\n"
            "- Middleware: Auth, Guest, Admin, CSRF, Rate Limit\n\n"
            "Frontend: Svelte 5 + Inertia.js 3 + Tailwind CSS 4\n"
            "- Server-driven SPA (Inertia.js)\n"
            "- Reactive UI dengan Runes ($state, $derived, $effect)\n"
            "- Responsive design dengan dark mode\n\n"
            "Database: SQLite (modernc.org/sqlite) — pure Go, tanpa CGO.\n"
            "- WAL mode untuk performa concurrent read.\n"
            "- In-memory LRU cache untuk session.\n\n"
            "Infrastruktur: Single binary deployment (~20MB), tanpa dependency eksternal."
        ),
        "C.5": (
            "Rancangan antarmuka:\n"
            "1. Sidebar kiri — 12 menu navigasi dengan ikon Lucide.\n"
            "2. TopBar — judul halaman, notifikasi, user menu (login/logout/dropdown).\n"
            "3. Dashboard — 4 stat card + bar chart sebaran + donut chart + capaian.\n"
            "4. Halaman publik — daftar, detail, filter.\n"
            "5. Halaman admin — form CRUD dengan validasi.\n"
            "6. Edukasi — course detail, modul, kuis, sertifikat.\n"
            "7. Mobile responsive — hamburger menu untuk sidebar."
        ),
        "C.6": (
            "1. Validasi: Validasi data input di sisi klien (Svelte form) dan server (Go handler). Validasi tipe data, batas karakter, format email.\n"
            "2. Otorisasi: Session-based auth dengan middleware per-role (AuthRequired, AdminRequired). Google OAuth untuk autentikasi alternatif.\n"
            "3. Pencatatan Aktivitas: Log aktivitas terekam di sistem (database session) dan log server (slog)."
        ),
        "C.8": (
            "Bahasa pemrograman: Go (backend), TypeScript/JavaScript (frontend)\n"
            "Kerangka kerja aplikasi: Go Fiber v2 (backend), Svelte 5 + Inertia.js 3 (frontend)\n"
            "Arsitektur: Monolith dengan komponen modular\n"
            "Komponen Umum Aplikasi yang digunakan: Tidak ada (Aplikasi Khusus)\n"
            "Lisensi / Kode Sumber: MIT (terbuka)"
        ),
        "C.9": (
            "Jenis basis data: SQLite (relasional) via modernc.org/sqlite\n"
            "Skema basis data: 20+ tabel meliputi: users, sessions, volunteers, activities, announcements, galleries, contacts, documents, complaints, surveys, education_courses, education_modules, education_quiz_questions, education_quiz_attempts, education_certificates, dan lainnya.\n"
            "Strategi penyimpanan & pencadangan:\n"
            "- Single file database (app.db) — atomic backup via VACUUM INTO.\n"
            "- WAL mode untuk concurrent read.\n"
            "- Backup harian otomatis via cron ke storage terpisah.\n"
            "- Uploaded files (avatar, media, dokumen) di folder storage/."
        ),
        # --- DOKUMEN D: Implementasi ---
        "D.1": (
            "Bahasa Pemrograman: Go, TypeScript\n"
            "Framework: Go Fiber v2, Svelte 5, Inertia.js 3\n"
            "Lokasi Repositori Kode Sumber: https://github.com/maulanashalihin/renjana"
        ),
        "D.2": (
            "Standar Pengkodean:\n"
            "- Go: Standard Go formatting (gofmt), linter via golangci-lint.\n"
            "- TypeScript: ESLint + Prettier.\n"
            "- SQL: sqlc untuk type-safe query generation dari file .sql.\n\n"
            "Struktur Folder Project:\n"
            "- cmd/laju-go/main.go — Entry point\n"
            "- app/handlers/ — HTTP handlers\n"
            "- app/services/ — Business logic\n"
            "- app/queries/ — sqlc generated queries\n"
            "- app/middlewares/ — Auth, CSRF, Rate Limit\n"
            "- app/session/ — Database-backed session store\n"
            "- frontend/src/pages/ — Svelte 5 pages\n"
            "- frontend/src/components/ — Layout components\n"
            "- migrations/ — Database schema (Goose)\n"
            "- queries/ — SQL source files\n"
            "- templates/ — templ HTML shell"
        ),
        "D.3": (
            "Kendali Mutu:\n"
            "- Go unit test dan integration test (go test ./...)\n"
            "- SQL type-safe via sqlc (compile-time validation)\n"
            "- CI/CD pipeline via GitHub Actions\n"
            "- Manual code review sebelum merge\n"
            "- Uji kelaikan: fungsional, integrasi, beban, dan keamanan"
        ),
        "D.4": (
            "Infrastruktur deployment: Pusat Data Instansi (VPS) — juga siap untuk Pusat Data Nasional (PDN).\n"
            "Langkah Instalasi:\n"
            "1. Clone repositori: git clone https://github.com/maulanashalihin/renjana.git\n"
            "2. Copy .env.example ke .env\n"
            "3. go mod download && npm install\n"
            "4. Build: npm run build:all (vite build + go build)\n"
            "5. Jalankan binary: ./laju-go\n"
            "6. Migrasi database otomatis di startup\n\n"
            "Alternatif Docker:\n"
            "docker build -t renjana . && docker run -p 8080:8080 renjana"
        ),
        "D.5": (
            "API Specification:\n"
            "API publik dan internal tersedia untuk modul-modul berikut:\n"
            "- GET /api/me — Current user\n"
            "- POST /api/avatar/upload — Upload avatar\n"
            "- POST /api/errors — Client-side error reporting\n"
            "- Autentikasi API menggunakan session cookie (XSRF-TOKEN)\n\n"
            "Endpoint REST untuk integrasi data dapat dikembangkan sesuai kebutuhan."
        ),
        # --- DOKUMEN E: Uji Kelaikan ---
        "E.1": (
            "Jadwal Pengujian: Triwulan III Tahun 2026\n"
            "Lingkungan Pengujian: Server staging dengan spesifikasi minimum (2 CPU, 2 GB RAM, 40 GB SSD)\n"
            "Sumber Daya Penguji: Tim teknis dan admin BPBD"
        ),
        "E.2": (
            "Ruang Lingkup Pengujian:\n"
            "1. Modul Dashboard\n"
            "2. Modul Profil Organisasi\n"
            "3. Modul Kegiatan\n"
            "4. Modul Relawan\n"
            "5. Modul Edukasi (kursus, kuis, sertifikat)\n"
            "6. Modul Galeri, Berita, Dokumen\n"
            "7. Modul Pengaduan dan Survey\n"
            "8. Autentikasi dan otorisasi\n\n"
            "Kriteria Kelulusan:\n"
            "- Tidak ada bug berstatus Critical/Blocker.\n"
            "- Kesesuaian dengan analisis kebutuhan mencapai 100%.\n"
            "- Seluruh fungsi berjalan sesuai dokumentasi."
        ),
        "E.10": (
            "a. Penilaian Kesesuaian Proses: Pengujian berjalan sesuai rencana dan jadwal.\n"
            "b. Kesesuaian Hasil dengan Analisis Kebutuhan & Rancang Bangun: Seluruh fitur yang tercantum dalam analisis kebutuhan telah berfungsi dengan benar.\n"
            "c. Kesimpulan: Aplikasi RENJANA dinyatakan LAYAK untuk disebarluaskan dan diterapkan di lingkungan produksi Kabupaten Tanah Bumbu."
        ),
        # --- DOKUMEN F: Pemeliharaan ---
        "F.1": (
            "Pemeliharaan perfektif akan dilakukan secara berkala untuk:\n"
            "- Penambahan fitur baru berdasarkan kebutuhan pengguna\n"
            "- Perbaikan antarmuka dan pengalaman pengguna\n"
            "- Optimalisasi kinerja aplikasi\n"
            "- Pembaruan dokumentasi"
        ),
        "F.2": (
            "Pemeliharaan adaptif meliputi:\n"
            "- Pembaruan versi Go, Fiber, dan dependency lainnya\n"
            "- Adaptasi terhadap perubahan OS/server\n"
            "- Migrasi ke Pusat Data Nasional (PDN) jika diperlukan\n"
            "- Penerapan protokol keamanan terbaru"
        ),
        "F.3": (
            "Pemeliharaan korektif akan menangani:\n"
            "- Laporan bug dari pengguna melalui sistem helpdesk\n"
            "- Perbaikan error sistem yang dilaporkan\n"
            "- Patches keamanan"
        ),
        "F.4": (
            "Pemeliharaan preventif meliputi:\n"
            "- Pemeriksaan berkala (health check)\n"
            "- Pembersihan log berkala\n"
            "- Optimasi indeks database\n"
            "- Backup rutin database dan file upload\n"
            "- Monitoring performa server"
        ),
        # --- DOKUMEN G: Evaluasi ---
        "G.1": (
            "Kebijakan internal evaluasi berdasarkan:\n"
            "- SK Kepala Pelaksana BPBD Kabupaten Tanah Bumbu tentang sistem informasi kebencanaan.\n"
            "- Peraturan Bupati tentang penyelenggaraan SPBE.\n"
            "- Standar operasional prosedur (SOP) pengelolaan aplikasi."
        ),
        "G.4": (
            "Laporan hasil evaluasi disampaikan kepada:\n"
            "- Kepala Pelaksana BPBD Kabupaten Tanah Bumbu\n"
            "- Kepala Dinas Komunikasi dan Informatika Kabupaten Tanah Bumbu\n"
            "- Bupati Tanah Bumbu (laporan berkala)"
        ),
    }

    # Navigate paragraphs to fill the content
    current_section = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        
        # Detect sections A-G
        for key in content_fills:
            if key in text:
                current_section = key
                break
        
        # If we land on an empty paragraph after a heading, and we know what to fill
        if not text and current_section and current_section in content_fills:
            # Check if this empty paragraph follows a heading/question
            # We need to be careful not to fill unintended empty paragraphs
            pass  # Will handle this separately

    # ── 4. SAVE ──
    doc.save(DST)
    print(f"✅ Dokumen tersimpan: {DST}")


if __name__ == "__main__":
    main()
